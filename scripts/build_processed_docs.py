#!/usr/bin/env python3
"""Build the curated lawyer deliverables into exports_processed/.

Unlike scripts/build_lawyer_docs.py (which assembles comments and memo entries
mechanically from the workspace database), the CONTENT here is hand-curated:
every margin comment and memo line was written editorially for counsel review.
The clause-election splicing that preserves the original Word formatting is
reused unchanged, and the spliced text is still verified against the app's own
resolver output.

Inputs:  exports/export_data.json (produced by scripts/build_export_data.mjs)
Outputs: exports_processed/Adjacent Capital - Term Sheet (Resolved Draft).docx
         exports_processed/Adjacent Capital - Decisions and Open Questions.docx
"""

from __future__ import annotations

import json
import re
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parent))
from docx_splice import align, apply_ops, cell_chars, rebuild_cell

ROOT = Path(__file__).resolve().parents[1]
EXPORTS = ROOT / "exports"
OUT_DIR = ROOT / "exports_processed"
ORIGINAL = ROOT / "sources" / "FORM - Venture Capital_Private Equity Fund Form Term Sheet.docx"

AUTHOR = "Adjacent Capital"
INITIALS = "AC"
FONT = "Times New Roman"

data = json.loads((EXPORTS / "export_data.json").read_text(encoding="utf-8"))
sections = data["sections"]


def clean(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


# ---------------------------------------------------------------------------
# Curated margin comments for the term sheet, keyed by clause.
# Clauses not listed get no comment. Rejected clauses are struck through and
# always get their curated removal note.
# ---------------------------------------------------------------------------

CLAUSE_COMMENTS: dict[str, str] = {
    "sec-03-investment-objectives-and-program": (
        "Please refine the strategy language. The mandate should cover frontier AI model developers and the "
        "surrounding stack (inference, infrastructure, and tooling), and we want maximal flexibility on the form "
        "of investment — direct equity, SPV interests, forwards, and other secondary-access structures "
        "(consider \u201cincluding but not limited to\u201d phrasing)."
    ),
    "sec-04-capital-commitments": (
        "The minimum commitment is shown as $5 million, but we are inclined to set it at $1 million for now, with "
        "the General Partner keeping discretion to accept smaller amounts. Please advise on what is typical for a "
        "fund of this size."
    ),
    "sec-08-management": (
        "Flagged for redraft once the management-entity question is settled — see the companion memo. We need to "
        "decide whether there is a single platform Manager LLC across funds or no separate management company, "
        "with management fees paid directly to the General Partner."
    ),
    "sec-09-sponsor-capital-commitment": (
        "Please remove. We would prefer no required sponsor commitment. If market practice or investor "
        "expectations effectively require one, please advise on the minimum credible level and whether it can be "
        "satisfied on a cashless basis (management fee waiver, promissory notes, or contributed securities)."
    ),
    "sec-11-management-fee-offset": (
        "Two questions on this clause: (1) do we need the second half of it at all, and (2) what is standard for "
        "the maximum aggregate management fee over the life of the fund?"
    ),
    "sec-12-limited-partner-advisory-committee": (
        "Do we need an LPAC at all for a fund of this profile? Please advise on standard practice. Note that "
        "several other provisions key off LPAC consent, so this decision ripples through the document."
    ),
    "sec-20-investment-limitations": (
        "Please restructure the concentration limits. We want the flexibility to invest up to the entire fund in "
        "a single position — for example, a Schedule A (Anthropic) with no concentration cap, and a Schedule B "
        "whitelist of companies capped at roughly 20% per issuer. Please also remove the prohibitions on weaponry "
        "and gambling businesses. Separately, we would like your advice on whether concentration limits need to "
        "be stated in the term sheet at all. The form's bracketed choices are left in place here since the clause "
        "is being restructured."
    ),
    "sec-23-warehoused-securities": (
        "Struck for now because we do not fully understand how this provision works in practice. Please walk us "
        "through it — if warehousing would give us useful flexibility for pre-closing purchases, we may want it "
        "back in."
    ),
    "sec-24-distributions": (
        "Confirming that we want a full General Partner catch-up in the waterfall."
    ),
    "sec-26-general-partner-clawback": (
        "Please remove. We would prefer no GP clawback — please advise whether institutional LPs will accept "
        "that, and if a clawback is effectively required, how the obligation is typically allocated among the "
        "ultimate carry recipients."
    ),
    "sec-29-other-competitive-activity-successor-fund": (
        "Please remove. We do not want restrictions on other activities or successor funds. Please advise on the "
        "likelihood of LP pushback and any standard middle-ground formulations."
    ),
}

# Small editorial text fixes applied to the spliced output. Each entry patches
# both the splice ops and the expected resolver text so verification still holds.
TEXT_FIXES: dict[str, list[tuple[str, str]]] = {
    # The form's blank had no unit suffix and the recorded value was just "500";
    # from context the cap is $500 million.
    "sec-04-capital-commitments": [("500", "500 million")],
}


def build_term_sheet() -> Path:
    doc = Document(str(ORIGINAL))
    table = doc.tables[0]

    # Title block: [FUND], LP -> Adjacent Capital, LP
    for para in doc.paragraphs[:6]:
        texts = [r.text for r in para.runs]
        if "FUND" in texts:
            for run in para.runs:
                if run.text == "[":
                    run.text = ""
                elif run.text == "FUND":
                    run.text = "Adjacent Capital"
                elif run.text == "]":
                    run.text = ""

    verify_failures: list[str] = []

    for entry in sections:
        if entry["isGroup"]:
            continue
        key = entry["stableKey"]
        row = table.rows[entry["row"]]
        title_cell, body_cell = row.cells[0], row.cells[1]
        status = entry["status"]

        if status == "rejected":
            for para in body_cell.paragraphs:
                for run in para.runs:
                    run.font.strike = True
                    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        elif status in ("accepted", "rewrite") and entry["fullyResolved"] and entry["ops"]:
            ops = entry["ops"]
            expected = entry["finalText"]
            for old, new in TEXT_FIXES.get(key, []):
                ops = [
                    {**op, "text": new} if op.get("text") == old else op
                    for op in ops
                ]
                expected = expected.replace(old, new, 1)
            chars = cell_chars(body_cell)
            positions = align(entry["originalBody"], chars)
            apply_ops(chars, positions, ops)
            rebuild_cell(body_cell, chars)
            if clean(body_cell.text) != expected:
                verify_failures.append(key)

        comment = CLAUSE_COMMENTS.get(key)
        if comment:
            anchor_runs = [r for p in title_cell.paragraphs for r in p.runs] or [
                r for p in body_cell.paragraphs for r in p.runs
            ]
            if anchor_runs:
                doc.add_comment(anchor_runs, text=comment, author=AUTHOR, initials=INITIALS)

    if verify_failures:
        raise SystemExit(f"spliced text does not match resolver output for: {verify_failures}")

    out = OUT_DIR / "Adjacent Capital - Term Sheet (Resolved Draft).docx"
    doc.save(out)
    return out


# ---------------------------------------------------------------------------
# Curated memo. Decisions first, then open questions. Anything handled by a
# term sheet margin comment is deliberately NOT repeated here.
# ---------------------------------------------------------------------------

# (label, statement, margin note for counsel or None)
DECISIONS: list[tuple[str, list[tuple[str, str, str | None]]]] = [
    (
        "Fund and entities",
        [
            ("Fund", "Adjacent Capital, LP, a Delaware limited partnership.", None),
            ("General Partner", "A fund-specific GP entity for each fund.", None),
            ("Sponsor HoldCo", "Yes — a sponsor holding company will sit above the Manager and GP entities.", None),
            (
                "Carry vehicle",
                "A separate CarryCo will receive the carried interest.",
                "Please confirm this is the standard approach, or whether carry can simply be allocated inside the "
                "GP LLC operating agreement without a separate entity.",
            ),
            (
                "Non-voting carry participants",
                "We want the ability to grant carry to non-voting economic participants (for example, someone who "
                "helps bring in capital) without giving them governance rights.",
                None,
            ),
        ],
    ),
    (
        "Regulatory and offering path",
        [
            (
                "Fund exemption",
                "Leaning toward Section 3(c)(7) with all investors as qualified purchasers.",
                "Please confirm this path works for our expected investor base.",
            ),
            (
                "Offering exemption",
                "Rule 506(b).",
                "Please confirm, and give us the practical pros and cons of 506(b) relationship rules versus "
                "506(c) verification.",
            ),
        ],
    ),
    (
        "Economics",
        [
            ("Fund size", "Target of approximately $100 million; hard cap of $500 million.", None),
            ("Management fee", "2.0% per annum on commitments, stepping down after the investment period (as elected in the term sheet).", None),
            ("Waterfall", "European (whole-of-fund) waterfall with a full GP catch-up.", None),
            (
                "Placement fees",
                "Placement costs should ideally be borne by the LPs rather than the sponsor.",
                "How is this best structured — can it be set aside from the LP pool or baked into an all-in price "
                "for the asset?",
            ),
        ],
    ),
    (
        "Investment mandate",
        [
            (
                "Whitelist approach",
                "The fund will invest in a Schedule A whitelist of named late-stage private technology companies.",
                None,
            ),
            (
                "Changing the whitelist",
                "Additions and removals will require limited partner consent; we need a workable structure for this.",
                None,
            ),
            (
                "Outside the whitelist",
                "Investments outside the whitelist will be permitted with LPAC or majority-LP approval.",
                None,
            ),
            (
                "Instruments",
                "The fund may invest through SPV interests, tender vehicles, forward contracts, contractual "
                "rights, and other secondary-access structures.",
                None,
            ),
            (
                "Non-U.S. exposure",
                "The fund should be able to invest in non-U.S. companies and non-U.S. holding structures.",
                "Does this raise legal or tax issues we should plan for?",
            ),
            (
                "Follow-ons",
                "Follow-on investments after the investment period should be permitted through the fifth "
                "anniversary of the initial closing with majority-in-interest consent.",
                "Please confirm the mechanics here reflect what is standard.",
            ),
        ],
    ),
    (
        "Capital calls",
        [
            (
                "Notice period",
                "Capital call notice of 3–5 business days, to support time-sensitive secondary purchases.",
                "Please confirm this is workable and what is standard.",
            ),
        ],
    ),
    (
        "Placement and sourcing",
        [
            (
                "Placement compensation",
                "We need the ability to compensate people for LP placement and for deal sourcing, potentially "
                "including carried interest.",
                "Can LP placement be paid out of management fees, and what should the structure look like?",
            ),
            (
                "Finders",
                "We would like to be able to use non-broker-dealer finders and consultants, ideally with carry as "
                "a component of compensation.",
                None,
            ),
            (
                "Sourcing arrangements",
                "We have proposed language permitting the Manager, GP, or affiliates to enter into sourcing, "
                "origination, consulting, brokerage, referral, or strategic relationship agreements in connection "
                "with investments.",
                "Is this language needed, and is it drafted correctly?",
            ),
            (
                "Offering materials",
                "Offering materials and investor communications will be approved by the management entity.",
                None,
            ),
            (
                "Disclosure",
                "We expect to disclose placement compensation to investors.",
                "Please advise on the standard scope and placement of that disclosure.",
            ),
        ],
    ),
]

# (question, sub-note or None)
OPEN_QUESTIONS: list[tuple[str, list[tuple[str, str | None]]]] = [
    (
        "Structure",
        [
            (
                "Management entity: we have gone back and forth between a single platform Manager LLC across all "
                "funds and having no separate management company at all, with management fees paid directly to the "
                "GP. What are the trade-offs, and what do you recommend for our situation?",
                "This decision also drives the redraft of the Management clause in the term sheet.",
            ),
            (
                "If we use a CarryCo, should it own the GP LLC, or merely receive carry economics from the GP?",
                None,
            ),
            (
                "We still need to finalize voting control at each level (GP, Manager, Sponsor HoldCo, CarryCo). "
                "Guidance on typical control structures would be helpful.",
                None,
            ),
            (
                "Are we better served by asset-specific funds (one fund per company) or a single fund with the "
                "whitelist mandate?",
                "This drives the Investment Objectives clause and the whitelist mechanics.",
            ),
        ],
    ),
    (
        "Capital calls and defaults",
        [
            ("What emergency or special capital call notice period is available for time-sensitive secondary deals?", None),
            ("Can the GP call capital for deposits, broken-deal expenses, legal diligence, and reserves?", None),
            ("What default remedies should apply if an LP misses a capital call?", None),
            ("Can the fund borrow under a bridge facility or subscription line, and can the GP or an affiliate advance funds to the fund?", None),
        ],
    ),
    (
        "Placement (buy-side)",
        [
            ("Who may solicit LPs, and must any compensated LP placement run through a registered broker-dealer?", None),
            ("Should placement compensation be paid by the Fund, the GP, the Manager, CarryCo, or an affiliate — and is it treated as a Fund Expense, a Manager expense, or GP/carry sharing?", None),
            ("Will placement parties receive a tail on LPs they introduced before termination of the arrangement?", None),
        ],
    ),
    (
        "Sourcing (sell-side)",
        [
            ("Who may source secondary opportunities, and who has authority to approve sourced deals?", None),
            ("What forms of compensation can deal sources receive (cash fees, deal-by-deal carry, fund carry, revenue share), and do such payments raise broker-dealer concerns?", None),
            ("Can a deal source negotiate terms or only make introductions — and can a deal source also represent the seller?", None),
            ("Are sourcing fees Fund Expenses or Manager/GP expenses?", None),
            ("What exclusivity, non-circumvention, or tail protections are standard for sourcing relationships, and what disclosures are required for affiliated or related-party sourcing?", None),
        ],
    ),
    (
        "Future funds, SPVs, and allocations",
        [
            ("What carveouts do we need for co-invest vehicles, single-deal SPVs, feeders, blockers, alternative investment vehicles, holding vehicles, and overflow vehicles?", None),
            ("Can the Manager run multiple funds at once, and how should investment opportunities be allocated among Fund I, later funds, co-invests, parallel vehicles, SPVs, affiliates, and insiders?", None),
            ("Can LPs receive co-investment rights, and can insiders invest personally alongside the fund?", None),
            ("Can the fund invest through a vehicle that charges its own fees or carry?", None),
        ],
    ),
    (
        "Regulatory, compliance, and operations",
        [
            ("Will the Manager rely on an investment adviser exemption or register, and what state-level adviser requirements apply?", None),
            ("Who should handle AML/KYC and bad-actor checks, and who will serve as fund administrator?", None),
            ("Who values private positions and SPV interests?", None),
            ("What insurance and indemnity coverage should be maintained, and what records must we keep for placement, sourcing, investment decisions, valuation, and conflicts?", None),
            ("How should transfer restrictions, ROFRs, company consents, tender rules, and lockups be disclosed to investors?", None),
        ],
    ),
    (
        "Fund lifecycle, distributions, and tax",
        [
            ("What should our side letter and MFN policy be?", None),
            ("How should subsequent-closing true-ups and interest work?", None),
            ("What recycling and reinvestment of proceeds should be permitted?", None),
            ("How should in-kind distributions of non-marketable private shares be handled?", None),
            ("What accommodations do tax-exempt and non-U.S. investors need?", None),
        ],
    ),
]


def _tnr(run, size=12, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    return run


def build_memo() -> Path:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(12)

    for text in ("ADJACENT CAPITAL, LP", "FUND FORMATION — DECISIONS AND OPEN QUESTIONS", f"{date.today():%B %d, %Y}"):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _tnr(para.add_run(text), bold=True)

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _tnr(
        intro.add_run(
            "Part 1 summarizes the decisions and working positions we have taken; margin comments flag the points "
            "where we would like your confirmation or advice. Part 2 collects the questions that remain open on "
            "our side. Clause-specific requests (deletions, rewrites, and drafting questions) are marked as "
            "comments in the companion document, \u201cAdjacent Capital \u2014 Term Sheet (Resolved Draft),\u201d "
            "and are not repeated here."
        ),
        italic=True,
    )

    def part_heading(text: str):
        para = doc.add_paragraph()
        _tnr(para.add_run(text.upper()), bold=True)

    def topic_heading(text: str):
        para = doc.add_paragraph()
        run = _tnr(para.add_run(text), bold=True)
        run.underline = True

    part_heading("Part 1 — Decisions and working positions")
    for topic, entries in DECISIONS:
        topic_heading(topic)
        for label, statement, note in entries:
            para = doc.add_paragraph(style="List Bullet")
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _tnr(para.add_run(f"{label}. "), bold=True)
            _tnr(para.add_run(statement))
            if note:
                doc.add_comment(para.runs, text=note, author=AUTHOR, initials=INITIALS)

    part_heading("Part 2 — Open questions")
    for topic, entries in OPEN_QUESTIONS:
        topic_heading(topic)
        for question, note in entries:
            para = doc.add_paragraph(style="List Bullet")
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _tnr(para.add_run(question))
            if note:
                note_para = doc.add_paragraph(style="List Bullet 2")
                _tnr(note_para.add_run(note), size=10, italic=True)

    out = OUT_DIR / "Adjacent Capital - Decisions and Open Questions.docx"
    doc.save(out)
    return out


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ts = build_term_sheet()
    memo = build_memo()
    print(f"Wrote {ts}")
    print(f"Wrote {memo}")


if __name__ == "__main__":
    main()
