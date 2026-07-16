#!/usr/bin/env python3
"""Build the two lawyer deliverables from exports/export_data.json.

Doc 1: Resolved term sheet - starts from the ORIGINAL Orrick .docx so the
        formatting (two-column table, Times New Roman, bold defined terms,
        numbered lists inside cells) is preserved exactly. Clause elections are
        spliced into the cells at run level; rejected clauses are struck
        through; notes/asks ride as margin comments anchored to clause titles.
Doc 2: Decisions and open questions memo, styled to match the term sheet
        (Times New Roman, centered caps title, justified body).

Both files open in Google Docs with the margin comments intact.
"""

from __future__ import annotations

import json
import re
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parent))
from docx_splice import align, apply_ops, cell_chars, rebuild_cell

ROOT = Path(__file__).resolve().parents[1]
EXPORTS = ROOT / "exports"
ORIGINAL = ROOT / "sources" / "FORM - Venture Capital_Private Equity Fund Form Term Sheet.docx"

AUTHOR = "Adjacent Capital"
INITIALS = "AC"
FONT = "Times New Roman"

data = json.loads((EXPORTS / "export_data.json").read_text(encoding="utf-8"))
sections = data["sections"]
issues = {row["stable_key"]: row for row in data["issues"]}


def issue_answer(key: str) -> str:
    return (issues[key].get("answer") or "").strip()


def issue_fu(key: str) -> str:
    return (issues[key].get("follow_up_notes") or "").strip()


def clean(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


# ---------------------------------------------------------------------------
# Document 1: resolved term sheet on top of the original file
# ---------------------------------------------------------------------------

def build_term_sheet() -> Path:
    doc = Document(str(ORIGINAL))
    table = doc.tables[0]

    # Title block: [FUND], LP -> Adjacent Capital, LP  (runs: '[','FUND',']',', ','LP')
    for para in doc.paragraphs[:6]:
        runs = para.runs
        texts = [r.text for r in runs]
        if "FUND" in texts:
            for run in runs:
                if run.text == "[":
                    run.text = ""
                elif run.text == "FUND":
                    run.text = "Adjacent Capital"
                elif run.text == "]":
                    run.text = ""

    verify_failures: list[str] = []

    for entry in sections:
        if entry["isGroup"]:
            continue
        row = table.rows[entry["row"]]
        title_cell, body_cell = row.cells[0], row.cells[1]
        status = entry["status"]
        comment_parts: list[str] = []

        if status == "rejected":
            for para in body_cell.paragraphs:
                for run in para.runs:
                    run.font.strike = True
                    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            comment_parts.append("REQUEST: Please remove this clause.")
            if entry["rewriteText"]:
                comment_parts.append(entry["rewriteText"])
            if entry["notes"]:
                comment_parts.append(entry["notes"])
        elif status in ("accepted", "rewrite"):
            if entry["fullyResolved"] and entry["ops"]:
                chars = cell_chars(body_cell)
                positions = align(entry["originalBody"], chars)
                apply_ops(chars, positions, entry["ops"])
                rebuild_cell(body_cell, chars)
                if clean(body_cell.text) != entry["finalText"]:
                    verify_failures.append(entry["stableKey"])
            elif not entry["fullyResolved"]:
                comment_parts.append(
                    f"Only {entry['electionsResolved']} of {entry['electionsTotal']} bracket choices were made "
                    "before we decided to restructure this clause, so the original form text is left as is."
                )
            if status == "rewrite":
                rewrite = entry["rewriteText"]
                if rewrite and rewrite.lower() not in (entry["finalText"] or "").lower():
                    comment_parts.append(f"REWRITE REQUESTED: {rewrite}")
                elif not rewrite:
                    comment_parts.append(
                        "Marked for rewrite in our review. Our elections are applied in the text - "
                        "please refine the drafting."
                    )
            if entry["notes"]:
                comment_parts.append(entry["notes"])
        # status == "none": plain clause, leave untouched

        if comment_parts:
            anchor_runs = [r for p in title_cell.paragraphs for r in p.runs] or [
                r for p in body_cell.paragraphs for r in p.runs
            ]
            if anchor_runs:
                doc.add_comment(anchor_runs, text="\n\n".join(comment_parts), author=AUTHOR, initials=INITIALS)

    if verify_failures:
        raise SystemExit(f"spliced text does not match resolver output for: {verify_failures}")

    out = EXPORTS / "Adjacent Capital - Term Sheet (Resolved Draft).docx"
    doc.save(out)
    return out


# ---------------------------------------------------------------------------
# Document 2: decisions and open questions (styled to match)
# ---------------------------------------------------------------------------

DECISIONS: list[tuple[str, list[tuple[str, str, str | None]]]] = [
    (
        "Structure and entities",
        [
            ("Fund entity", issue_answer("q-001-will-the-fund-be-a-delaware-lp-or-delaware-llc"), None),
            ("GP entity", issue_answer("q-002-will-the-gp-be-fund-specific-for-each-fund"), None),
            (
                "Sponsor HoldCo above the Manager and GP entities",
                issue_answer("q-004-will-there-be-a-sponsor-holdco-above-the-manager-and-gp-entities"),
                None,
            ),
            (
                "Carry vehicle",
                "Separate CarryCo. "
                + issue_answer("q-005-do-we-need-a-fund-specific-carryco-sponsorco-or-can-carry-be-split-in-"),
                "Please confirm whether a separate carry company is standard here, or whether carry can simply be "
                "split inside the GP LLC operating agreement.",
            ),
            (
                "Non-voting carry participants",
                issue_answer("q-008-can-non-voting-economic-participants-receive-carry-without-governance-"),
                None,
            ),
        ],
    ),
    (
        "Regulatory and offering path",
        [
            (
                "Fund exemption",
                issue_answer("decision-3-fund-exemption-path"),
                "Please confirm the 3(c)(7) / all-QP path.",
            ),
            (
                "Offering exemption",
                issue_answer("decision-4-offering-path"),
                issue_fu("decision-4-offering-path") or None,
            ),
        ],
    ),
    (
        "Economics",
        [
            (
                "Waterfall",
                issue_answer("q-017-european-waterfall-or-deal-by-deal-waterfall"),
                None,
            ),
            ("GP catch-up", "Yes (recorded on the Distributions clause).", None),
            (
                "Placement fees / fee offsets",
                issue_answer("q-014-will-placement-fees-organization-expenses-transaction-fees-or-monitori"),
                issue_fu("q-014-will-placement-fees-organization-expenses-transaction-fees-or-monitori") or None,
            ),
        ],
    ),
    (
        "Investment mandate",
        [
            (
                "Whitelist approach",
                "Yes — the fund will be limited to a Schedule A whitelist of companies. "
                + issue_answer("q-020-what-counts-as-a-late-stage-private-technology-company"),
                None,
            ),
            (
                "Changing the whitelist",
                issue_answer("q-022-who-can-add-or-remove-whitelist-companies-and-with-what-approval"),
                None,
            ),
            (
                "Investing outside the whitelist",
                "Yes, with LPAC or majority LP approval.",
                None,
            ),
            (
                "SPV interests, tender vehicles, forwards, secondary-access vehicles",
                "Yes. "
                + issue_answer("change-04-restriction-on-blind-pool-fund-investments-carveouts-for-weird-vehicle"),
                issue_fu("change-04-restriction-on-blind-pool-fund-investments-carveouts-for-weird-vehicle") or None,
            ),
            (
                "Non-U.S. companies / holding structures",
                issue_answer("q-025-can-the-fund-invest-in-non-u-s-companies-or-non-u-s-holding-structures"),
                issue_fu("q-025-can-the-fund-invest-in-non-u-s-companies-or-non-u-s-holding-structures") or None,
            ),
            (
                "Follow-on investments after the investment period",
                issue_answer("q-027-will-the-fund-allow-follow-on-investments-after-the-investment-period"),
                "Recorded as a partial answer — please confirm the intended follow-on mechanics.",
            ),
        ],
    ),
    (
        "Capital calls",
        [
            (
                "Capital call timing",
                issue_answer("decision-5-capital-call-design"),
                issue_fu("decision-5-capital-call-design") or None,
            ),
        ],
    ),
    (
        "Placement and sourcing",
        [
            (
                "Conditional carry / incentive compensation",
                issue_answer("change-01-ability-to-give-conditional-carry-or-incentive-payments"),
                "Open follow-up: can LP placement be paid out of management fees, and what does the structure look like?",
            ),
            (
                "Non-broker-dealer finders",
                issue_answer("q-042-can-non-bd-finders-or-consultants-be-used-and-under-what-restrictions"),
                None,
            ),
            (
                "Sourcing / origination arrangements (requested language)",
                issue_answer("change-02-sourcing-origination-and-strategic-relationship-arrangements"),
                "Do we need this language at all?",
            ),
            (
                "Offering materials approval",
                issue_answer("q-048-who-approves-offering-materials-and-investor-communications"),
                None,
            ),
            (
                "Placement compensation disclosure",
                issue_answer("decision-7-placement-compensation-disclosure"),
                "Please advise on standard practice for disclosing placement compensation.",
            ),
        ],
    ),
]

OPEN_QUESTIONS: list[tuple[str, list[tuple[str, str | None]]]] = [
    (
        "Structure and entities",
        [
            (
                "Manager entity — conflicting answers recorded: one answer says a single platform Manager LLC "
                "across all funds (with the Management LLC approving offering materials); another says no separate "
                "management company, with management fees paid to the GPs. Which structure should we use?",
                None,
            ),
            ("If CarryCo is used, does it own the GP LLC or merely receive carry economics from the GP?", None),
            (
                "How are clawback obligations allocated among ultimate carry recipients?",
                "Related: we rejected the General Partner Clawback clause in the term sheet — please advise whether "
                "that is viable and how clawback should work if not.",
            ),
            ("Who has voting control at each level: GP, Manager, Sponsor HoldCo, CarryCo?", "Partially answered ('we') — needs completion."),
        ],
    ),
    (
        "Investment mandate",
        [
            (
                "Are we doing asset-specific funds, or one fund with a generic whitelist?",
                "This decision drives the Investment Objectives clause and the whitelist mechanics.",
            ),
            ("How should transfer restrictions, ROFRs, company consent, tender rules, and lockups be disclosed?", None),
        ],
    ),
    (
        "Capital calls and warehousing",
        [
            ("What is the emergency/special capital call notice period for time-sensitive secondary deals?", None),
            ("Can the GP call capital for deposits, broken-deal expenses, legal diligence, and reserves?", None),
            ("What default remedies apply if an LP misses a capital call?", None),
            ("Can the fund borrow under a bridge facility or subscription line?", None),
            ("Can the GP or an affiliate advance funds to the fund?", None),
            (
                "What approvals are needed for warehoused investments, and how are warehousing costs, interest, "
                "legal fees, and conflict disclosures handled?",
                "Related: we struck the Warehoused Securities clause because we did not understand it well enough — "
                "please walk us through whether we need it.",
            ),
        ],
    ),
    (
        "Buy-side placement / capital formation",
        [
            ("Who can solicit LPs?", None),
            ("Must any compensated LP placement be through a registered broker-dealer?", None),
            ("Is placement compensation paid by the Fund, GP, Manager, CarryCo, or an affiliate?", None),
            ("Is placement compensation treated as a Fund Expense, Manager expense, or GP/carry sharing?", None),
            ("Will placement parties receive a tail on LPs introduced before termination?", None),
            ("How will 506(b) relationship rules or 506(c) verification rules be handled?", "Please cover pros/cons of each path."),
        ],
    ),
    (
        "Sell-side sourcing / origination",
        [
            ("Who can source secondary opportunities, and who has authority to approve sourced deals?", None),
            ("Can sourcing parties receive cash fees, deal-by-deal carry, fund carry, or revenue share?", None),
            ("Can a deal source negotiate terms, or only introduce opportunities? Can a deal source also represent sellers?", None),
            ("Do deal-source payments create broker-dealer concerns?", None),
            ("Are sourcing fees Fund Expenses or Manager/GP expenses?", None),
            ("Are there exclusivity, non-circumvention, or tail periods for sourcing relationships?", None),
            ("What disclosures are required for affiliated or related-party sourcing?", None),
            ("How are conflicts handled if a source also participates in another vehicle?", None),
        ],
    ),
    (
        "Future funds, SPVs, and allocations",
        [
            (
                "What carveouts are needed for co-invest vehicles, single-deal SPVs, feeders, blockers, alternative "
                "investment vehicles, holding vehicles, and overflow vehicles?",
                "Related: we struck the Other Competitive Activity / Successor Fund clause — we do not want a "
                "successor-fund restriction. Please advise on implications.",
            ),
            ("Can the Manager manage multiple funds at once?", None),
            ("How are opportunities allocated among Fund I, Fund II, co-invests, parallel funds, SPVs, affiliates, and insiders?", None),
            ("Can LPs get co-investment rights?", None),
            ("Can insiders invest personally alongside the fund?", None),
            ("Can the fund invest through a vehicle that charges its own fees/carry?", None),
            (
                "Will LPAC approval be needed for certain conflicts or allocation decisions?",
                "Related: we are unsure whether we need an LPAC at all — please advise on what is standard.",
            ),
        ],
    ),
    (
        "Fees and sponsor economics",
        [
            ("What is standard for the maximum management fee?", "Also: do we need the second part of the Management Fee Offset clause?"),
            (
                "Do we need a GP/sponsor capital commitment at all? Ideally zero.",
                "We struck the Sponsor Capital Commitment clause; if some commitment is required, can it be cashless "
                "(notes, warehoused securities, or fee waiver)?",
            ),
        ],
    ),
    (
        "Regulatory, compliance, and operations",
        [
            ("Will the manager rely on an investment adviser exemption or register?", None),
            ("What state-level adviser requirements apply?", None),
            ("Who handles AML/KYC and bad actor checks?", None),
            ("Who will be fund administrator?", None),
            ("Who values private positions and SPV interests?", None),
            ("What cyber/privacy/data-room procedures are needed?", None),
            ("What insurance and indemnity coverage should be maintained?", None),
            ("What records must be kept for placement, sourcing, investment decisions, valuation, and conflicts?", None),
        ],
    ),
    (
        "Fund lifecycle, distributions, and tax",
        [
            ("Side letter and MFN policy?", None),
            ("Subsequent-closing true-up and interest?", None),
            ("Recycling and reinvestment of proceeds?", None),
            ("In-kind distributions of non-marketable private shares?", None),
            ("Tax-exempt and non-U.S. investor accommodations?", None),
        ],
    ),
    (
        "Term sheet mark-up (see comments in the resolved draft)",
        [
            (
                "Investment Limitations: we want the ability to potentially invest the entire fund in one asset — "
                "possibly a Schedule A that is just Anthropic and a Schedule B list of companies capped at 20% each. "
                "We also want the ability to invest in weaponry and gambling. Please restructure.",
                None,
            ),
            ("Management clause: rewrite requested (structure depends on the Manager-entity question above).", None),
            ("Investment Objectives: rewrite requested — strategy language for frontier model development and inference and the associated stack.", None),
        ],
    ),
]


def _tnr(run, size=12, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    return run


def build_memo() -> Path:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(12)

    for text in ("ADJACENT CAPITAL, LP", "FUND FORMATION — DECISIONS AND OPEN QUESTIONS", f"{date.today():%B %d, %Y}"):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _tnr(para.add_run(text), bold=True)

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _tnr(
        intro.add_run(
            "Part 1 lists the decisions and working positions we have taken, with follow-ups for counsel in the "
            "margin. Part 2 is the consolidated list of open questions. Companion document: “Adjacent Capital — "
            "Term Sheet (Resolved Draft),” which carries clause-level elections and mark-up comments."
        ),
        italic=True,
    )

    def part_heading(text: str):
        para = doc.add_paragraph()
        _tnr(para.add_run(text.upper()), bold=True)

    def topic_heading(text: str):
        para = doc.add_paragraph()
        run = _tnr(para.add_run(text), bold=True)
        run.underline = True

    part_heading("Part 1 — Decisions and working positions")
    for topic, entries in DECISIONS:
        topic_heading(topic)
        for label, answer, follow_up in entries:
            para = doc.add_paragraph(style="List Bullet")
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _tnr(para.add_run(f"{label}: "), bold=True)
            _tnr(para.add_run(answer if answer else "(recorded without notes)"))
            if follow_up:
                doc.add_comment(para.runs, text=follow_up, author=AUTHOR, initials=INITIALS)

    part_heading("Part 2 — Open questions")
    intro2 = doc.add_paragraph()
    intro2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _tnr(
        intro2.add_run(
            "These remain open on our side. Questions that were effectively answered elsewhere have been removed; "
            "where an open question connects to a change we made in the term sheet, the connection is noted."
        ),
        italic=True,
    )

    for topic, entries in OPEN_QUESTIONS:
        topic_heading(topic)
        for question, note in entries:
            para = doc.add_paragraph(style="List Bullet")
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _tnr(para.add_run(question))
            if note:
                note_para = doc.add_paragraph(style="List Bullet 2")
                _tnr(note_para.add_run(note), size=10, italic=True)

    out = EXPORTS / "Adjacent Capital - Decisions and Open Questions.docx"
    doc.save(out)
    return out


def main() -> None:
    ts = build_term_sheet()
    memo = build_memo()
    print(f"Wrote {ts}")
    print(f"Wrote {memo}")


if __name__ == "__main__":
    main()
