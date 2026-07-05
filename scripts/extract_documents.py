#!/usr/bin/env python3
"""Extract the Orrick term sheet and open-items memo into app seed data."""

from __future__ import annotations

import json
import re
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "sources"
DATA_DIR = ROOT / "data"

TERM_SHEET = SOURCE_DIR / "FORM - Venture Capital_Private Equity Fund Form Term Sheet.docx"
OPEN_ITEMS = SOURCE_DIR / "blind_pool_fund_open_items_and_drafting_changes.docx"
CHATGPT_SHARE_URL = "https://chatgpt.com/share/6a47ea7d-00d8-83e8-af2e-4c84dc4c580d"


def clean(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def slugify(text: str, fallback: str = "item") -> str:
    slug = re.sub(r"[^a-z0-9]+", "-", text.lower()).strip("-")
    return slug[:70] or fallback


def unique(values: Iterable[str]) -> list[str]:
    seen: set[str] = set()
    out: list[str] = []
    for value in values:
        if value and value not in seen:
            seen.add(value)
            out.append(value)
    return out


def read_paragraphs(path: Path) -> list[str]:
    doc = Document(str(path))
    return [clean(p.text) for p in doc.paragraphs if clean(p.text)]


def read_tables(path: Path) -> list[list[list[str]]]:
    doc = Document(str(path))
    tables: list[list[list[str]]] = []
    for table in doc.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([clean(cell.text) for cell in row.cells])
        tables.append(rows)
    return tables


def extract_term_sections(path: Path) -> list[dict]:
    doc = Document(str(path))
    sections: list[dict] = []
    current_group = "Front Matter"
    for row_index, row in enumerate(doc.tables[0].rows):
        title = clean(row.cells[0].text)
        body = clean(row.cells[1].text)
        is_group = not body or body == title
        section_id = f"sec-{row_index:02d}-{slugify(title, 'section')}"
        if is_group:
            current_group = title
        sections.append(
            {
                "id": section_id,
                "row": row_index,
                "title": title,
                "body": body,
                "group": title if is_group else current_group,
                "isGroup": is_group,
            }
        )
    return sections


def build_title_index(sections: list[dict]) -> dict[str, str]:
    index: dict[str, str] = {}
    for section in sections:
        if section["isGroup"]:
            continue
        index[section["title"].lower()] = section["id"]
    for section in sections:
        index.setdefault(section["title"].lower(), section["id"])
    return index


def anchor_ids(index: dict[str, str], titles: list[str]) -> list[str]:
    return unique(index.get(title.lower(), "") for title in titles)


# Short alias so the curated maps below stay readable.
HOLDING_VEHICLES = "Holding Vehicles; Feeder Vehicle; Alternative Investment Vehicles; Parallel Funds"
SUCCESSOR_FUND = "Other Competitive Activity; Successor Fund"


# Curated, question-level anchoring. Each memo question maps to the specific
# term-sheet clauses its answer actually changes, plus a priority and a short
# "how to decide" note. Keyed by the exact question text in the memo so a
# regenerated extraction stays aligned with the source document.
QUESTION_CURATION: dict[str, dict] = {
    # A. Entity structure and ownership
    "Will the fund be a Delaware LP or Delaware LLC?": {
        "sections": ["The Fund", "The Fund Agreement"],
        "priority": "high",
        "note": "The Orrick form assumes a Delaware LP governed by an LPA with a GP LLC. Choosing an LLC fund instead changes governance, tax boilerplate, and every Partner/Partnership reference, so settle this before any other drafting.",
    },
    "Will the GP be fund-specific for each fund?": {
        "sections": ["Management", "The Fund"],
        "priority": "medium",
        "note": "A fund-specific GP is market standard: it isolates liability and carry per vintage and simplifies Fund II. A shared GP saves one entity but entangles the funds' economics and liabilities.",
    },
    "Will there be a single platform Manager LLC across all funds?": {
        "sections": ["Management", "Management Fee"],
        "priority": "medium",
        "note": "A single platform Manager is the usual pattern: it is the registered/exempt adviser and collects the management fee from each fund under an investment management agreement.",
    },
    "Will there be a Sponsor HoldCo above the Manager and GP entities?": {
        "sections": ["Management"],
        "priority": "low",
        "note": "A HoldCo mainly helps if outside owners or several principals need platform-level economics. For a two-principal Fund I it usually just adds an entity.",
    },
    "Do we need a fund-specific CarryCo/SponsorCo, or can carry be split in the GP LLC operating agreement?": {
        "sections": ["Distributions", "General Partner Clawback", "Management"],
        "priority": "high",
        "note": "Carry can be split inside the GP LLC operating agreement. A separate CarryCo is cleaner when non-voting participants, conditional/vested awards, or clawback sharing among recipients are expected - all likely here given placement and sourcing carry.",
    },
    "If CarryCo is used, does it own the GP LLC or merely receive carry economics from the GP?": {
        "sections": ["Management", "Distributions"],
        "priority": "medium",
        "note": "Economics-only participation keeps control with the principals; ownership of the GP moves governance too. Most sponsors keep CarryCo as an economics vehicle only.",
    },
    "Who has voting control at each level: GP, Manager, Sponsor HoldCo, CarryCo?": {
        "sections": ["Management", "Removal of the General Partner"],
        "priority": "medium",
        "note": "Map control before granting anyone economics. LPs will also care: the removal-of-GP clause interacts with who actually controls the GP entity.",
    },
    "Can non-voting economic participants receive carry without governance rights?": {
        "sections": ["Distributions", "Management"],
        "priority": "medium",
        "note": "Yes, if the GP LLC or CarryCo agreement creates a non-voting economic class. Confirm securities-law treatment of granting those interests, especially to placement parties.",
    },
    "How are clawback obligations allocated among ultimate carry recipients?": {
        "sections": ["General Partner Clawback", "Distributions", "Limited Partner Giveback"],
        "priority": "high",
        "note": "The Fund-level clawback sits on the GP; the GP/CarryCo agreement must push it through to every carry recipient (escrow, holdback, or several guarantee), or the principals end up covering departed participants' share. Also settle the mechanics: the form tests the clawback only at wind-up, but if the waterfall is deal-by-deal LPs will push for interim clawback tests, an after-tax calculation, and escrow or personal guarantees from carry recipients.",
    },
    # B. Fund economics
    "What is the target fund size and hard cap?": {
        "sections": ["Capital Commitments", "Closings"],
        "priority": "high",
        "note": "Fills the blanks in Capital Commitments and drives the 3(c)(1) LP-count math, minimum commitment, and fee projections.",
    },
    "What is the minimum LP commitment, and can the GP waive it?": {
        "sections": ["Capital Commitments"],
        "priority": "medium",
        "note": "The form already gives the GP discretion to accept smaller subscriptions; just set the stated minimum.",
    },
    "What is the management fee rate and base: committed capital, invested capital, NAV, or stepped-down rate?": {
        "sections": ["Management Fee"],
        "priority": "high",
        "note": "The form offers committed-capital fee with post-investment-period stepdown options. For a whitelist secondaries fund, LPs may push for invested-capital or stepped-down fees.",
    },
    "Does the management fee begin at initial closing or only when capital is drawn?": {
        "sections": ["Management Fee", "Closings"],
        "priority": "medium",
        "note": "Fee from initial closing is standard but interacts with how much capital is drawn upfront; also decide the subsequent-closing fee true-up.",
    },
    "Will placement fees, organization expenses, transaction fees, or monitoring fees offset management fees?": {
        "sections": ["Management Fee Offset", "Management Fee", "Fund Expenses"],
        "priority": "medium",
        "note": "The offset section is optional in the form. Decide offsets now because they determine whether placement and sourcing costs are ultimately borne by LPs or by the Manager. Pair this with a cap on organizational and offering expenses borne by the fund (commonly $250k-$1M depending on size, excess borne by the Manager) - the form currently charges formation costs to the fund without a stated cap.",
    },
    "What is the carry percentage?": {
        "sections": ["Distributions"],
        "priority": "high",
        "note": "Sets the waterfall's carried-interest split (typically 20%; secondaries funds sometimes vary). Everything in the CarryCo/participation discussion keys off this number.",
    },
    "Is there a preferred return or hurdle?": {
        "sections": ["Distributions"],
        "priority": "high",
        "note": "Decide hurdle (commonly 8%) or no-hurdle. For a fast-recycling secondaries strategy the hurdle plus catch-up mechanics materially change GP economics.",
    },
    "European waterfall or deal-by-deal waterfall?": {
        "sections": ["Distributions", "General Partner Clawback"],
        "priority": "high",
        "note": "Deal-by-deal pays carry earlier but makes the clawback provision do real work; whole-fund (European) is more LP-friendly and simpler to administer.",
    },
    "Will there be GP catch-up?": {
        "sections": ["Distributions"],
        "priority": "medium",
        "note": "If there is a preferred return, decide full, partial, or no catch-up - this is a pure economics negotiation point.",
    },
    "What is the GP commitment, and can it be cashless, notes, warehoused securities, or fee waiver?": {
        "sections": ["Sponsor Capital Commitment", "Warehoused Securities", "Capital Contributions"],
        "priority": "high",
        "note": "The form's Sponsor Capital Commitment already brackets cash, cashless, warehoused securities, and notes. Confirm the percentage (1-2% is common) and which funding forms LPs will accept.",
    },
    # C. Investment mandate and whitelist
    "What counts as a late-stage private technology company?": {
        "sections": ["Investment Objectives and Program"],
        "priority": "high",
        "note": "Define objectively (e.g., minimum valuation, revenue stage, or named-list approach) so compliance with the mandate is checkable at deal time.",
    },
    "Will the fund be limited to a Schedule A whitelist of companies?": {
        "sections": ["Investment Objectives and Program", "Investment Limitations"],
        "priority": "high",
        "note": "A whitelist makes the blind pool much easier to market but constrains execution; pair it with a defined process to amend the schedule.",
    },
    "Who can add or remove whitelist companies, and with what approval?": {
        "sections": ["Investment Limitations", "Limited Partner Advisory Committee", "Amendments"],
        "priority": "medium",
        "note": "Typical design: GP proposes, LPAC approves additions; removals unrestricted. Avoid requiring full LP amendment consent for routine list changes.",
    },
    "Can the fund invest in companies outside the whitelist with LPAC or majority LP approval?": {
        "sections": ["Investment Limitations", "Limited Partner Advisory Committee"],
        "priority": "medium",
        "note": "An LPAC-approved exception path preserves flexibility without an amendment; decide the approval threshold and any basket size.",
    },
    "Can the fund invest in SPV interests, tender vehicles, forward contracts, contractual rights, or secondary-access vehicles?": {
        "sections": ["Investment Objectives and Program", "Investment Limitations", HOLDING_VEHICLES],
        "priority": "high",
        "note": "Core to the strategy: the mandate and the investment limitations must both expressly permit indirect exposure instruments, or every secondary deal risks a consent process.",
    },
    "Can the fund invest in non-U.S. companies or non-U.S. holding structures?": {
        "sections": ["Investment Objectives and Program", HOLDING_VEHICLES, "Taxation"],
        "priority": "medium",
        "note": "If yes, the vehicles section should permit blockers/feeders for tax reasons and the PPM needs non-U.S. risk and withholding disclosure.",
    },
    "What are single-company concentration limits?": {
        "sections": ["Investment Limitations"],
        "priority": "medium",
        "note": "Typical range is 15-25% of commitments per issuer. A whitelist strategy may justify a higher cap; look-through SPV exposure should count toward it.",
    },
    "Will the fund allow follow-on investments after the investment period?": {
        "sections": ["Investment Period", "Investment Limitations"],
        "priority": "medium",
        "note": "The Investment Period section should carve out follow-ons and pending transactions; decide any cap on post-period follow-on capital.",
    },
    "Will there be sector exclusions or ESG/prohibited activity restrictions?": {
        "sections": ["Investment Limitations"],
        "priority": "low",
        "note": "Often handled in side letters rather than fund-wide restrictions; decide the baseline now so side-letter asks are manageable.",
    },
    "How should transfer restrictions, ROFRs, company consent, tender rules, and lockups be disclosed?": {
        "sections": ["General Risks", "Investment Objectives and Program", "Valuation"],
        "priority": "medium",
        "note": "These are portfolio-level execution risks unique to secondaries: disclose in the PPM risk factors and reflect in the valuation policy, since ROFRs and consent rights can break signed deals.",
    },
    # D. Capital calls, speed, and warehousing
    "What percentage, if any, is drawn at closing?": {
        "sections": ["Capital Contributions", "Closings"],
        "priority": "medium",
        "note": "The form brackets an optional initial drawdown. An upfront 10-25% draw plus reserves is the simplest way to fund fast secondary closings.",
    },
    "What is the standard capital call notice period?": {
        "sections": ["Capital Contributions"],
        "priority": "high",
        "note": "The form offers 10/15/20 business days; secondary allocations often close faster than that. Consider 10 business days standard with a shorter expedited path.",
    },
    "What is the emergency/special capital call notice period for time-sensitive secondary deals?": {
        "sections": ["Capital Contributions"],
        "priority": "high",
        "note": "Add an expedited call (e.g., 3-5 business days) for identified transactions, paired with the bridge facility and GP-advance mechanics as a backstop.",
    },
    "Can the GP call capital for deposits, broken-deal expenses, legal diligence, and reserves?": {
        "sections": ["Capital Contributions", "Fund Expenses"],
        "priority": "medium",
        "note": "Make sure the call mechanics and the Fund Expenses definition both cover deposits and broken-deal costs, which are routine in competitive secondary processes.",
    },
    "What default remedies apply if an LP misses a capital call?": {
        "sections": ["Failure to Make Capital Contributions", "Capital Contributions"],
        "priority": "medium",
        "note": "The form leaves remedies to the fund agreement. Standard menu: interest on late amounts, dilution/forfeiture up to 50%, forced transfer, and loss of voting rights.",
    },
    "Can the fund borrow under a bridge facility or subscription line?": {
        "sections": ["Borrowing", HOLDING_VEHICLES],
        "priority": "medium",
        "note": "The form permits borrowing subject to a bracketed cap tied to commitments or uncalled commitments. Pick the cap (commonly 20-30% of commitments) and decide whether guarantees and leverage at SPV or holding-vehicle level count against it - with deals executed through SPVs, a fund-level-only cap is meaningless. Fund borrowing can also create UBTI for tax-exempt LPs.",
    },
    "Can the GP or affiliate advance funds to the fund?": {
        "sections": ["Borrowing", "General Partner Expenses"],
        "priority": "medium",
        "note": "GP advances are the fastest bridge for a closing deadline. Authorize them expressly, with interest at cost and repayment from the next call to avoid a conflicts fight later.",
    },
    "Can the GP or affiliate warehouse securities and later sell them to the fund?": {
        "sections": ["Warehoused Securities", "Sponsor Capital Commitment"],
        "priority": "high",
        "note": "The form's optional Warehoused Securities clause permits transfers at cost plus expenses. Keep it - warehousing is core to executing before the fund closes - and decide whether warehoused positions can satisfy the GP commitment.",
    },
    "What approvals are needed for warehoused investments?": {
        "sections": ["Warehoused Securities", "Limited Partner Advisory Committee", "Investment Limitations"],
        "priority": "medium",
        "note": "Affiliate sales to the fund are classic conflict territory: decide whether cost-plus-expenses transfers within a set window are pre-authorized or need LPAC sign-off.",
    },
    "How are warehousing costs, interest, legal fees, and conflict disclosures handled?": {
        "sections": ["Warehoused Securities", "Fund Expenses"],
        "priority": "medium",
        "note": "Define the transfer price formula (cost plus carrying costs at a stated rate) and disclose the arrangement in the PPM conflicts section.",
    },
    # E. Buy-side placement / capital formation
    "Who can solicit LPs?": {
        "sections": ["Investor Qualifications", "General Risks"],
        "priority": "high",
        "note": "Principals and Manager personnel can solicit under the issuer exemption; anyone compensated for raising capital raises broker-dealer issues. Inventory the expected placement parties before signing anything.",
    },
    "Must any compensated LP placement be through a registered broker-dealer?": {
        "sections": ["Investor Qualifications"],
        "priority": "high",
        "note": "Transaction-based compensation to an unregistered finder is the classic foot-fault; the safe default is BD-only for success-fee arrangements.",
    },
    "Can non-BD finders or consultants be used, and under what restrictions?": {
        "sections": ["Investor Qualifications"],
        "priority": "medium",
        "note": "Only with flat-fee/consulting structures that avoid transaction-based pay, no negotiation role, and counsel sign-off. Document the analysis per relationship.",
    },
    "Can placement compensation include cash fees, management fee sharing, carried interest, revenue share, or expense reimbursement?": {
        "sections": ["Management Fee Offset", "Fund Expenses", "Distributions"],
        "priority": "medium",
        "note": "Carry-sharing is usually the cleanest for non-BD parties because it is an internal GP/CarryCo split rather than a fund-paid fee, but it still needs conflicts disclosure and BD analysis if tied to commitments raised.",
    },
    "Is compensation paid by the Fund, GP, Manager, CarryCo, or affiliate?": {
        "sections": ["Fund Expenses", "General Partner Expenses", "Management Fee Offset"],
        "priority": "medium",
        "note": "Who pays determines who bears the cost: Fund-paid placement fees are LP money (and often trigger a fee offset); Manager/GP-paid comes out of sponsor economics.",
    },
    "Is the compensation treated as Fund Expense, Manager expense, or GP/carry sharing?": {
        "sections": ["Fund Expenses", "Management Fee Offset"],
        "priority": "medium",
        "note": "Keep the term sheet, LPA expense section, and offset provision consistent - LPs read all three, and mismatches invite negotiation.",
    },
    "Will placement-linked carry require disclosure in the PPM and/or term sheet?": {
        "sections": ["General Risks", "Side Letters"],
        "priority": "medium",
        "note": "Disclose the category of arrangement (carry sharing with placement parties) in the PPM conflicts/fees sections; specific splits can stay in the separate agreements.",
    },
    "Will placement parties receive a tail on LPs introduced before termination?": {
        "sections": ["Fund Expenses"],
        "priority": "low",
        "note": "A tail is a placement-agreement term, not a fund-document term; decide the standard tail (commonly 6-12 months) for the form agreement.",
    },
    "Who approves offering materials and investor communications?": {
        "sections": ["Management", "Investor Qualifications"],
        "priority": "low",
        "note": "The Manager should have sole authority over offering materials, with placement parties contractually barred from creating their own.",
    },
    "How will 506(b) relationship rules or 506(c) verification rules be handled?": {
        "sections": ["Investor Qualifications", "General Risks"],
        "priority": "high",
        "note": "506(b) requires pre-existing relationships and no general solicitation; 506(c) allows solicitation but requires accreditation verification. Placement-party outreach practices must match whichever path is chosen.",
    },
    # F. Sell-side sourcing / origination
    "Who can source secondary opportunities?": {
        "sections": ["Management", "Investment Objectives and Program"],
        "priority": "medium",
        "note": "Anyone can introduce deals; the design question is who may be compensated and how. Keep investment authority with the GP regardless of source.",
    },
    "Can sourcing parties receive cash fees, deal-by-deal carry, fund carry, or revenue share?": {
        "sections": ["Fund Expenses", "Distributions", "Management Fee Offset"],
        "priority": "high",
        "note": "Deal-by-deal carry to a source is a fund-economics event that needs waterfall or GP-split plumbing; cash success fees raise BD issues if tied to securities transactions.",
    },
    "Can a deal source negotiate terms, or only introduce opportunities?": {
        "sections": ["Management"],
        "priority": "medium",
        "note": "Introduction-only keeps the source on the safer side of the broker-dealer line; negotiation authority plus success comp looks like brokerage.",
    },
    "Can a deal source also represent sellers?": {
        "sections": [SUCCESSOR_FUND, "Exculpation and Indemnification"],
        "priority": "medium",
        "note": "Dual representation is a live conflict: require disclosure, and decide whether LPAC consent is needed when a compensated source is on both sides.",
    },
    "Do deal-source payments create broker-dealer concerns?": {
        "sections": ["General Risks"],
        "priority": "high",
        "note": "Yes if compensation is transaction-based and the source effects securities transactions. This is a counsel sign-off item for each sourcing agreement, not a one-time decision.",
    },
    "Are sourcing fees Fund Expenses or Manager/GP expenses?": {
        "sections": ["Fund Expenses", "General Partner Expenses"],
        "priority": "medium",
        "note": "The form's GP Expenses section brackets whether investment-sourcing costs sit with the sponsor. Decide the default and mirror it in the sourcing form agreement.",
    },
    "Are there exclusivity, non-circumvention, or tail periods?": {
        "sections": [SUCCESSOR_FUND],
        "priority": "low",
        "note": "These live in the sourcing agreement; set house-standard positions (no exclusivity, narrow non-circumvention, short tail) for the template.",
    },
    "What disclosures are required for affiliated or related-party sourcing?": {
        "sections": ["Investment Limitations", "General Risks"],
        "priority": "medium",
        "note": "Affiliate-sourced deals should be disclosed in the PPM conflicts section and may need LPAC consent under the investment limitations.",
    },
    "How are conflicts handled if a source also participates in another vehicle?": {
        "sections": [SUCCESSOR_FUND, "Co-Investment Opportunities"],
        "priority": "medium",
        "note": "Combine allocation-policy treatment with LPAC disclosure: a source with its own SPV competing for the same allocation is the main real-world scenario.",
    },
    "Who has authority to approve sourced deals?": {
        "sections": ["Management", "Limited Partner Advisory Committee"],
        "priority": "medium",
        "note": "The GP's investment committee approves all deals; sources never hold approval rights. Only conflicted deals escalate to the LPAC.",
    },
    # G. Future funds, SPVs, and allocations
    "When can a successor fund be launched?": {
        "sections": [SUCCESSOR_FUND],
        "priority": "high",
        "note": "The form restricts successor blind pools during the investment period until a deployment threshold (often 70-75% invested/reserved). Pick the threshold and confirm SPVs don't count as successor funds.",
    },
    "What counts as a prohibited substantially similar blind-pool fund?": {
        "sections": [SUCCESSOR_FUND],
        "priority": "medium",
        "note": "Define narrowly: a multi-investor blind pool with a substantially overlapping mandate. Single-deal SPVs, co-invests, and feeders should be expressly excluded.",
    },
    "What carveouts are needed for co-invest vehicles, single-deal SPVs, feeders, blockers, alternative investment vehicles, holding vehicles, and overflow vehicles?": {
        "sections": [SUCCESSOR_FUND, HOLDING_VEHICLES, "Co-Investment Opportunities"],
        "priority": "high",
        "note": "List every execution-vehicle type as an express carveout from the successor-fund restriction; this is the memo's core drafting request and the biggest execution risk if missed.",
    },
    "Can the Manager manage multiple funds at once?": {
        "sections": [SUCCESSOR_FUND, "Management"],
        "priority": "medium",
        "note": "The form already permits other competitive activity; make sure the Manager's other mandates only trip the restriction if they are substantially similar blind pools.",
    },
    "How are opportunities allocated among Fund I, Fund II, co-invests, parallel funds, SPVs, affiliates, and insiders?": {
        "sections": [SUCCESSOR_FUND, "Co-Investment Opportunities"],
        "priority": "high",
        "note": "Adopt a written allocation policy (fund-first with capacity-based overflow to co-invest/SPVs is typical) and reference it in the PPM; regulators and LPs both ask for it.",
    },
    "Can LPs get co-investment rights?": {
        "sections": ["Co-Investment Opportunities", "Side Letters"],
        "priority": "medium",
        "note": "The form keeps co-invest offers discretionary. Hold that line in the LPA and grant any priority rights through side letters so MFN exposure is controlled.",
    },
    "Can insiders invest personally alongside the fund?": {
        "sections": ["Co-Investment Opportunities", "Investment Limitations"],
        "priority": "medium",
        "note": "Permit only after the fund takes its full allocation, with conflicts disclosure; unconstrained insider co-investment is a common LP diligence objection.",
    },
    "Can the fund invest through a vehicle that charges its own fees/carry?": {
        "sections": ["Investment Limitations", HOLDING_VEHICLES],
        "priority": "medium",
        "note": "Secondary access often requires buying into fee-bearing SPVs. Permit it with a cap or LPAC consent, and disclose the fee layering; a flat prohibition would block core strategy deals.",
    },
    "Will LPAC approval be needed for certain conflicts or allocation decisions?": {
        "sections": ["Limited Partner Advisory Committee", "Investment Limitations"],
        "priority": "medium",
        "note": "Reserve LPAC approval for genuine conflicts (affiliate transactions, warehousing outside pre-agreed terms, allocation exceptions) and keep routine execution GP-discretionary. Settle the LPAC's composition at the same time: the form brackets member count and appointment mechanics, and since the LPAC will carry real weight here (whitelist changes, warehousing, conflicts), keep it small (3-5 GP-appointed members from the largest LPs) and define its approval list precisely.",
    },
    # H. Regulatory, compliance, and operations
    "Will the manager rely on an investment adviser exemption or register?": {
        "sections": ["Management", "General Risks"],
        "priority": "high",
        "note": "Most sub-$150M private fund managers use the SEC private fund adviser exemption (ERA status) plus state rules. The answer drives compliance obligations across the platform.",
    },
    "What state-level adviser requirements apply?": {
        "sections": ["Management"],
        "priority": "medium",
        "note": "Check the manager's home state: some states require investment adviser registration even for exempt-reporting advisers.",
    },
    "Who handles AML/KYC and bad actor checks?": {
        "sections": ["Investor Qualifications"],
        "priority": "medium",
        "note": "Usually the fund administrator runs AML/KYC on subscriptions; Rule 506(d) bad-actor checks must also cover placement parties and beneficial owners of the GP.",
    },
    "Who is fund administrator?": {
        "sections": ["Reports", "Fund Expenses"],
        "priority": "medium",
        "note": "Select early - the administrator drives capital-call logistics, reporting cadence, and AML/KYC workflow, and admin fees are a named Fund Expense.",
    },
    "Will annual financials be audited?": {
        "sections": ["Reports"],
        "priority": "medium",
        "note": "The form promises annual financials; institutional LPs and the custody rule (if registered) effectively require an audit. Budget it as a Fund Expense.",
    },
    "What reporting cadence is promised to LPs?": {
        "sections": ["Reports"],
        "priority": "medium",
        "note": "The form's quarterly unaudited/annual package is standard; do not promise more in side letters than the administrator can produce.",
    },
    "Who values private positions and SPV interests?": {
        "sections": ["Valuation", "In-Kind Distributions"],
        "priority": "high",
        "note": "The GP values assets in good faith under the form, but stale secondary marks are a known regulator focus; adopt a written valuation policy covering indirect SPV positions and in-kind distributions.",
    },
    "What cyber/privacy/data-room procedures are needed?": {
        "sections": ["Reports"],
        "priority": "low",
        "note": "Operational item: data-room access controls and investor-data handling belong in the compliance manual, not the fund documents.",
    },
    "What insurance and indemnity coverage should be maintained?": {
        "sections": ["Exculpation and Indemnification", "Fund Expenses"],
        "priority": "medium",
        "note": "GP/Manager D&O and E&O coverage backstops the fund indemnity; the form lets the fund bear the premium as a Fund Expense. Confirm the indemnity's conduct standard too - the form exculpates absent material misconduct (vs a gross-negligence standard) - and whether covered persons extend to sourcing or placement parties, which LPs will resist.",
    },
    "What records must be kept for placement, sourcing, investment decisions, valuation, and conflicts?": {
        "sections": ["Reports", "Management"],
        "priority": "low",
        "note": "Keep contemporaneous files for solicitation, allocation, valuation, and conflicts decisions - these are the first documents requested in an exam or LP dispute.",
    },
}


# Questions the memo does not ask but the term sheet requires answers to.
# Sourced from a section-by-section gap review of the form term sheet.
GAP_QUESTIONS: list[dict] = [
    {
        "category": "I. Fund lifecycle and governance (gap review)",
        "title": "How long is the investment period, and what ends it early",
        "prompt": "How long is the investment period, and what events end it early?",
        "sections": ["Investment Period", "Management Fee"],
        "priority": "high",
        "note": "The form brackets the anniversary and early-termination triggers (key person event, GP removal, LP vote). The choice also controls the management-fee stepdown date. Secondaries funds often use a shorter period (2-3 years) than the venture default.",
    },
    {
        "category": "I. Fund lifecycle and governance (gap review)",
        "title": "What is the fund term, and who approves extensions",
        "prompt": "What is the fund term, and who approves extensions (GP discretion, LPAC, or LP consent)?",
        "sections": ["Term", "Dissolution"],
        "priority": "high",
        "note": "The form brackets term length and one-or-two-year extensions with GP/LPAC/LP approval options. Late-stage secondaries may exit faster than venture: consider an 8-year term with two one-year extensions.",
    },
    {
        "category": "I. Fund lifecycle and governance (gap review)",
        "title": "Who are the key persons and what happens on a key person event",
        "prompt": "Who are the key persons, what time commitment do they promise, and what happens on a key person event?",
        "sections": ["Time Commitment; Key Person Event", "Investment Period"],
        "priority": "high",
        "note": "The form suspends or terminates the investment period on a key person event. With a two-principal sponsor, decide whether losing either principal (or only both) triggers it, and whether LPs can vote to restart.",
    },
    {
        "category": "I. Fund lifecycle and governance (gap review)",
        "title": "On what standard can LPs remove the GP, and what happens to carry",
        "prompt": "Under what standard and LP vote can the GP be removed, and what happens to carry and the GP commitment on removal?",
        "sections": ["Removal of the General Partner", "Distributions"],
        "priority": "medium",
        "note": "The form allows for-cause removal with a bracketed carry reduction (often 20-50%). Decide the cause standard, the vote threshold, and whether there is any no-fault removal right.",
    },
    {
        "category": "I. Fund lifecycle and governance (gap review)",
        "title": "Amendment consent thresholds",
        "prompt": "What LP consent threshold applies to amendments, and which amendments require special or affected-partner consent?",
        "sections": ["Amendments", "Limited Partner Advisory Committee"],
        "priority": "medium",
        "note": "The form protects LPs from adverse amendments without consent. Confirm the general threshold (majority in interest is typical) and carve out administrative amendments the GP can make alone - important given the expected vehicle flexibility changes.",
    },
    {
        "category": "I. Fund lifecycle and governance (gap review)",
        "title": "Side letter and MFN policy",
        "prompt": "What side letter policy and MFN (most favored nation) rights will the fund offer?",
        "sections": ["Side Letters", "Investor Qualifications"],
        "priority": "medium",
        "note": "The form permits side letters including better economics. Since placement arrangements and co-invest rights will live in side letters, decide MFN scope early (commonly size-tiered) so one LP's deal doesn't propagate to everyone.",
    },
    {
        "category": "I. Fund lifecycle and governance (gap review)",
        "title": "Subsequent-closing true-up and interest",
        "prompt": "Will subsequent-closing investors pay a true-up plus interest on their share of earlier capital calls and fees?",
        "sections": ["Closings", "Capital Contributions"],
        "priority": "medium",
        "note": "The form permits later closings within a bracketed window. With early warehoused deals, later investors will be buying into appreciated positions - decide the true-up interest rate and whether late LPs share in pre-closing deals at cost.",
    },
    {
        "category": "J. Distributions, liability, and tax (gap review)",
        "title": "Recycling and reinvestment of proceeds",
        "prompt": "May the fund recycle or reinvest disposition proceeds, and within what limits?",
        "sections": ["Distributions", "Investment Period"],
        "priority": "high",
        "note": "The memo never addresses recycling, but a secondaries fund with quick flips needs it: without reinvestment rights, early exits shrink deployable capital. Typical formulation permits recycling proceeds received during the investment period up to 100-120% of commitments.",
    },
    {
        "category": "J. Distributions, liability, and tax (gap review)",
        "title": "In-kind distributions of non-marketable private shares",
        "prompt": "Can the fund distribute non-marketable private shares in kind, and how are they valued and transferred given company ROFRs and consent rights?",
        "sections": ["In-Kind Distributions", "Valuation", "Dissolution"],
        "priority": "high",
        "note": "The form only allows non-marketable in-kind distributions at dissolution. For a fund holding private secondaries, end-of-life positions may be unsaleable - the mechanics (valuation, ROFR compliance, LP election to decline) deserve real attention.",
    },
    {
        "category": "J. Distributions, liability, and tax (gap review)",
        "title": "LP giveback cap and time limit",
        "prompt": "What cap and time limit apply to LP givebacks for indemnification and other fund liabilities?",
        "sections": ["Limited Partner Giveback", "Exculpation and Indemnification"],
        "priority": "medium",
        "note": "The form allows the GP to recall distributions for liabilities. Market practice caps givebacks (commonly 25% of distributions or commitments) and sunsets them 2-3 years after distribution; the form leaves this to the LPA.",
    },
    {
        "category": "J. Distributions, liability, and tax (gap review)",
        "title": "Tax-exempt and non-U.S. investor accommodations",
        "prompt": "Will the fund accept tax-exempt or non-U.S. LPs, and are blockers or feeders needed for UBTI/ECI concerns?",
        "sections": ["Taxation", "ERISA", HOLDING_VEHICLES],
        "priority": "medium",
        "note": "Fund-level borrowing (subscription line, bridge) can create UBTI for tax-exempt LPs; non-U.S. LPs care about ECI. Decide whether to offer a blocker/feeder or restrict the investor base, and stay under the 25% ERISA plan-asset threshold.",
    },
]


# Curated anchors for the memo's immediate-decision table (keyed by decision label).
DECISION_CURATION: dict[str, list[str]] = {
    "Use separate CarryCo?": ["Distributions", "General Partner Clawback", "Management", "Sponsor Capital Commitment"],
    "Who owns Manager LLC?": ["Management", "Management Fee"],
    "Fund exemption path?": ["Investor Qualifications"],
    "Offering path?": ["Investor Qualifications", "General Risks"],
    "Capital-call design?": ["Capital Contributions", "Borrowing", "Warehoused Securities"],
    "Vehicle flexibility?": ["Investment Objectives and Program", "Investment Limitations", HOLDING_VEHICLES],
    "Placement compensation disclosure?": ["Management Fee Offset", "Fund Expenses", "Side Letters", "General Risks"],
    "Successor fund restriction?": [SUCCESSOR_FUND, "Co-Investment Opportunities"],
}


# Curated anchors for the requested drafting changes (keyed by memo heading).
CHANGE_CURATION: dict[str, list[str]] = {
    "A. Ability to give conditional carry or incentive payments": ["Distributions", "Management Fee Offset", "Fund Expenses", "General Partner Clawback"],
    "B. Sourcing, origination, and strategic relationship arrangements": ["Management", "Fund Expenses", SUCCESSOR_FUND, "Exculpation and Indemnification"],
    "C. Investment objectives and permitted investments": ["Investment Objectives and Program", "Investment Limitations"],
    "D. Restriction on blind-pool fund investments - carveouts for weird vehicles": ["Investment Limitations", HOLDING_VEHICLES],
    "E. Capital call timing and fast execution mechanics": ["Capital Contributions", "Borrowing", "Warehoused Securities", "Failure to Make Capital Contributions"],
    "F. Successor funds, parallel funds, co-invests, and SPV carveouts": [SUCCESSOR_FUND, "Co-Investment Opportunities", HOLDING_VEHICLES],
}


# Curated anchors for supporting documents (keyed by document name in the memo tables).
SUPPORTING_DOC_CURATION: dict[str, list[str]] = {
    "Fund LPA / Fund Agreement": ["The Fund Agreement"],
    "Subscription Agreement": ["Investor Qualifications", "Closings"],
    "PPM / Offering Memorandum": ["General Risks", "Investment Objectives and Program"],
    "Form D and state blue-sky filings": ["Investor Qualifications"],
    "Side Letter Template": ["Side Letters"],
    "GP LLC Operating Agreement": ["Management", "Distributions", "General Partner Clawback"],
    "Manager LLC Operating Agreement": ["Management", "Management Fee"],
    "Investment Management Agreement": ["Management", "Management Fee"],
    "Sponsor HoldCo Agreement, if used": ["Management"],
    "Fund-specific CarryCo / SponsorCo Agreement, if used": ["Distributions", "General Partner Clawback", "Management"],
    "Registered Broker-Dealer Placement Agreement": ["Investor Qualifications", "Fund Expenses"],
    "Generic Capital Formation / Placement Agreement": ["Investor Qualifications", "Fund Expenses", "Management Fee Offset"],
    "Conditional Carry Participation Agreement": ["Distributions", "General Partner Clawback"],
    "Finder / Consultant Agreement": ["Investor Qualifications", "General Risks"],
    "Sell-Side Sourcing / Origination Agreement": ["Fund Expenses", SUCCESSOR_FUND],
    "Secondary Transaction / Broker Agreement": ["Investment Objectives and Program", "Transfers of Interests and Withdrawals"],
    "Investment Policy / Whitelist Schedule": ["Investment Objectives and Program", "Investment Limitations"],
    "Allocation Policy": ["Co-Investment Opportunities", SUCCESSOR_FUND],
    "Capital Call Policy": ["Capital Contributions"],
    "Warehousing Policy": ["Warehoused Securities", "Limited Partner Advisory Committee"],
    "Conflict Policy": ["Investment Limitations", "Limited Partner Advisory Committee", SUCCESSOR_FUND],
    "Valuation Policy": ["Valuation", "In-Kind Distributions"],
    "Compliance Policies": ["Investor Qualifications", "Reports"],
}


def make_issue(
    *,
    issue_id: str,
    issue_type: str,
    title: str,
    prompt: str,
    source: str,
    term_section_ids: list[str],
    category: str = "",
    details: str = "",
    provisional_answer: str = "",
    priority: str = "medium",
    tags: list[str] | None = None,
) -> dict:
    return {
        "id": issue_id,
        "issueType": issue_type,
        "status": "open",
        "priority": priority,
        "category": category,
        "title": title,
        "prompt": prompt,
        "details": details,
        "provisionalAnswer": provisional_answer,
        "source": source,
        "termSectionIds": term_section_ids,
        "tags": tags or [],
    }


def extract_detailed_questions(paragraphs: list[str], index: dict[str, str]) -> list[dict]:
    issues: list[dict] = []
    current_category = ""
    in_section = False
    question_index = 1
    uncurated: list[str] = []
    for text in paragraphs:
        if text == "3. Detailed Questions to Answer":
            in_section = True
            continue
        if text == "4. Requested Changes / Drafting Instructions for Orrick":
            break
        if not in_section:
            continue
        if re.match(r"^[A-H]\.\s+", text):
            current_category = text
            continue
        if "?" not in text:
            continue
        title = text.rstrip("?")
        curation = QUESTION_CURATION.get(text)
        if curation is None:
            uncurated.append(text)
            curation = {"sections": [], "priority": "medium", "note": ""}
        issues.append(
            make_issue(
                issue_id=f"q-{question_index:03d}-{slugify(title)}",
                issue_type="question",
                title=title,
                prompt=text,
                source="Open-items memo - Section 3 detailed questions",
                category=current_category,
                details=curation["note"],
                term_section_ids=anchor_ids(index, curation["sections"]),
                priority=curation["priority"],
                tags=[slugify(current_category.split(".", 1)[-1].strip())],
            )
        )
        question_index += 1
    if uncurated:
        print(f"WARNING: {len(uncurated)} memo question(s) missing curation:")
        for text in uncurated:
            print(f"  - {text}")
    return issues


def build_gap_questions(index: dict[str, str]) -> list[dict]:
    issues: list[dict] = []
    for gap_index, gap in enumerate(GAP_QUESTIONS, start=1):
        issues.append(
            make_issue(
                issue_id=f"gap-{gap_index:03d}-{slugify(gap['title'])}",
                issue_type="question",
                title=gap["title"],
                prompt=gap["prompt"],
                source="Term sheet gap review - not covered in the open-items memo",
                category=gap["category"],
                details=gap["note"],
                term_section_ids=anchor_ids(index, gap["sections"]),
                priority=gap["priority"],
                tags=["term-sheet-gap"],
            )
        )
    return issues


def extract_decisions(tables: list[list[list[str]]], index: dict[str, str]) -> list[dict]:
    issues: list[dict] = []
    decision_table = tables[14]
    for row in decision_table[1:]:
        number, decision, option_a, option_b, comment = row[:5]
        prompt = f"{decision}: choose between the listed approaches or draft a custom answer."
        provisional = f"Option A: {option_a}\nOption B: {option_b}\nComment: {comment}"
        issues.append(
            make_issue(
                issue_id=f"decision-{number}-{slugify(decision)}",
                issue_type="decision",
                title=decision,
                prompt=prompt,
                source="Open-items memo - Section 5 immediate decisions",
                category="Immediate decisions",
                details=comment,
                provisional_answer=provisional,
                term_section_ids=anchor_ids(index, DECISION_CURATION.get(decision, [])),
                priority="high",
                tags=["immediate-decision"],
            )
        )
    return issues


# Note: the memo's Section 7 checklist is intentionally NOT extracted as issues.
# Every checklist line duplicates an immediate decision, a detailed question, or a
# supporting document that already exists in the queue; extracting it doubled up
# the workload list without adding information.


def extract_drafting_changes(paragraphs: list[str], tables: list[list[list[str]]], index: dict[str, str]) -> list[dict]:
    heads = [
        "A. Ability to give conditional carry or incentive payments",
        "B. Sourcing, origination, and strategic relationship arrangements",
        "C. Investment objectives and permitted investments",
        "D. Restriction on blind-pool fund investments - carveouts for weird vehicles",
        "E. Capital call timing and fast execution mechanics",
        "F. Successor funds, parallel funds, co-invests, and SPV carveouts",
    ]
    table_indices = [8, 9, 10, 11, 12, 13]
    instruction_prefixes = {
        heads[0]: "Add a general authorization/disclosure provision.",
        heads[1]: "Add a general authorization/disclosure provision.",
        heads[2]: "Revise the generic investment objective to cover these categories.",
        heads[3]: "Modify the blind-pool or fee-bearing vehicle restriction to preserve execution flexibility.",
        heads[4]: "Revise capital-call and funding mechanics to allow faster execution.",
        heads[5]: "Revise the successor-fund restriction and pair it with allocation rules.",
    }
    notes_by_head: dict[str, str] = {}
    for idx, text in enumerate(paragraphs):
        if text in heads:
            for nxt in paragraphs[idx + 1 : idx + 8]:
                if nxt.startswith("Counsel notes / decisions:"):
                    notes_by_head[text] = nxt
                    break

    issues: list[dict] = []
    for change_index, (head, table_index) in enumerate(zip(heads, table_indices), start=1):
        provision = tables[table_index][0][0]
        prompt = instruction_prefixes[head]
        details = notes_by_head.get(head, "")
        title = head.split(". ", 1)[1]
        issues.append(
            make_issue(
                issue_id=f"change-{change_index:02d}-{slugify(title)}",
                issue_type="drafting-change",
                title=title,
                prompt=prompt,
                source="Open-items memo - Section 4 requested drafting changes",
                category="Requested drafting changes",
                details=details,
                provisional_answer=provision,
                term_section_ids=anchor_ids(index, CHANGE_CURATION.get(head, [])),
                priority="high",
                tags=["drafting-change"],
            )
        )
    return issues


def extract_supporting_documents(tables: list[list[list[str]]], index: dict[str, str]) -> list[dict]:
    issues: list[dict] = []
    for table_index in [2, 3, 4, 5]:
        header = tables[table_index][0][0]
        for row_index, row in enumerate(tables[table_index][1:], start=1):
            name, purpose, notes = row[:3]
            title = f"{header}: {name}"
            issues.append(
                make_issue(
                    issue_id=f"doc-{table_index}-{row_index}-{slugify(name)}",
                    issue_type="supporting-document",
                    title=title,
                    prompt=purpose,
                    source="Open-items memo - Section 1 additional documents needed",
                    category=f"Supporting documents: {header.lower()}",
                    details=notes,
                    term_section_ids=anchor_ids(index, SUPPORTING_DOC_CURATION.get(name, [])),
                    priority="low",
                    tags=["supporting-document", slugify(header)],
                )
            )
    return issues


def main() -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    term_sections = extract_term_sections(TERM_SHEET)
    index = build_title_index(term_sections)
    memo_paragraphs = read_paragraphs(OPEN_ITEMS)
    memo_tables = read_tables(OPEN_ITEMS)

    issues = []
    issues.extend(extract_decisions(memo_tables, index))
    issues.extend(extract_drafting_changes(memo_paragraphs, memo_tables, index))
    issues.extend(extract_detailed_questions(memo_paragraphs, index))
    issues.extend(build_gap_questions(index))
    issues.extend(extract_supporting_documents(memo_tables, index))

    payload = {
        "meta": {
            "project": "Blind Pool Fund Term Sheet Workspace",
            "generatedAt": datetime.now().isoformat(timespec="seconds"),
            "sourceFiles": [
                {
                    "label": "Orrick venture capital / private equity fund form term sheet",
                    "path": str(TERM_SHEET),
                },
                {
                    "label": "Blind pool fund open items and drafting changes",
                    "path": str(OPEN_ITEMS),
                },
                {
                    "label": "Shared ChatGPT conversation",
                    "url": CHATGPT_SHARE_URL,
                    "note": "The public share page was available, but the clean transcript was not directly exposed through static fetch; the attached open-items memo was used as the structured seed.",
                },
            ],
        },
        "termSheet": {
            "title": "[FUND], LP - Summary of Principal Terms",
            "sections": term_sections,
        },
        "openItemsMemo": {
            "paragraphs": memo_paragraphs,
            "tables": memo_tables,
        },
        "issues": issues,
    }

    json_text = json.dumps(payload, indent=2, ensure_ascii=False)
    (DATA_DIR / "seed-data.json").write_text(json_text + "\n", encoding="utf-8")
    (DATA_DIR / "seed-data.js").write_text(
        "window.ORRICK_SEED_DATA = " + json_text + ";\n", encoding="utf-8"
    )

    print(f"Wrote {DATA_DIR / 'seed-data.json'}")
    print(f"Wrote {DATA_DIR / 'seed-data.js'}")
    print(f"Sections: {len(term_sections)}")
    print(f"Issues: {len(issues)}")


if __name__ == "__main__":
    main()
